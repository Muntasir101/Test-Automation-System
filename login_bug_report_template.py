from docx import Document
from docx.shared import Inches


class BugReportTemplate:
    def __init__(self):
        self.template = {
            "title": "Bug Report",
            "date": "Date",
            "reporter": "Reported by",
            "description": "Description",
            "steps_to_reproduce": "Steps to Reproduce",
            "expected_result": "Expected Result",
            "actual_result": "Actual Result",
            "severity": "Severity",
            "priority": "Priority"
        }

    def generate_report(self, details, screenshot_path, report_path):
        doc = Document()
        doc.add_heading(self.template["title"], 0)

        doc.add_heading(self.template["date"], level=1)
        doc.add_paragraph(details['date'])

        doc.add_heading(self.template["reporter"], level=1)
        doc.add_paragraph(details['reporter'])

        doc.add_heading(self.template["description"], level=1)
        doc.add_paragraph(details['description'])

        doc.add_heading(self.template["steps_to_reproduce"], level=1)
        doc.add_paragraph(details['step1'])
        doc.add_paragraph(details['step2'])
        doc.add_paragraph(details['step3'])

        doc.add_heading(self.template["expected_result"], level=1)
        doc.add_paragraph(details['expected_result'])

        doc.add_heading(self.template["actual_result"], level=1)
        doc.add_paragraph(details['actual_result'])

        doc.add_heading(self.template["severity"], level=1)
        doc.add_paragraph(details['severity'])

        doc.add_heading(self.template["priority"], level=1)
        doc.add_paragraph(details['priority'])

        if screenshot_path:
            doc.add_heading("Screenshot", level=1)
            doc.add_picture(screenshot_path, width=Inches(5))

        doc.save(report_path)
