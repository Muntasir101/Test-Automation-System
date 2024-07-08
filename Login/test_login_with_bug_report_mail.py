import pytest
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import os
import datetime
from docx import Document
from docx.shared import Inches
import yagmail


class BugReportTemplate:
    def __init__(self):
        self.template = {
            "title": "Bug Report",
            "date": "Date",
            "reporter": "Reported by",
            "description": "Description",
            "steps_to_reproduce": "Steps to Reproduce",
            "expected_result": "Expected Result",
            "actual_result": "Actual Result",
            "severity": "Severity",
            "priority": "Priority"
        }

    def generate_report(self, details, screenshot_path, report_path):
        doc = Document()
        doc.add_heading(self.template["title"], 0)

        doc.add_heading(self.template["date"], level=1)
        doc.add_paragraph(details['date'])

        doc.add_heading(self.template["reporter"], level=1)
        doc.add_paragraph(details['reporter'])

        doc.add_heading(self.template["description"], level=1)
        doc.add_paragraph(details['description'])

        doc.add_heading(self.template["steps_to_reproduce"], level=1)
        doc.add_paragraph(details['step1'])
        doc.add_paragraph(details['step2'])
        doc.add_paragraph(details['step3'])

        doc.add_heading(self.template["expected_result"], level=1)
        doc.add_paragraph(details['expected_result'])

        doc.add_heading(self.template["actual_result"], level=1)
        doc.add_paragraph(details['actual_result'])

        doc.add_heading(self.template["severity"], level=1)
        doc.add_paragraph(details['severity'])

        doc.add_heading(self.template["priority"], level=1)
        doc.add_paragraph(details['priority'])

        if screenshot_path:
            doc.add_heading("Screenshot", level=1)
            doc.add_picture(screenshot_path, width=Inches(5))

        doc.save(report_path)


def send_email(report_path, to_email):
    yag = yagmail.SMTP('test@gmail.com', 'nnsw mynz') # password must generated app password
    subject = "Automated Bug Report"
    contents = "Please find the attached bug report."
    yag.send(to_email, subject, contents, attachments=[report_path])
    print(f"Bug report sent to {to_email}")


@pytest.fixture
def driver():
    driver = webdriver.Firefox()
    driver.get('https://tutorialsninja.com/demo/index.php?route=account/login')
    yield driver
    driver.quit()


def capture_screenshot(driver, test_name):
    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    screenshot_name = f"{test_name}_{timestamp}.png"
    screenshot_path = os.path.join(os.path.dirname(__file__), 'failed_screenshots', screenshot_name)
    os.makedirs(os.path.dirname(screenshot_path), exist_ok=True)
    driver.save_screenshot(screenshot_path)
    return screenshot_path


def login(driver, email, password):
    wait = WebDriverWait(driver, 10)
    wait.until(EC.presence_of_element_located((By.NAME, 'email')))

    username_field = driver.find_element(By.NAME, 'email')
    username_field.clear()
    username_field.send_keys(email)

    password_field = driver.find_element(By.NAME, 'password')
    password_field.clear()
    password_field.send_keys(password)

    login_button = driver.find_element(By.CSS_SELECTOR, "[action] .btn-primary")
    login_button.click()


def generate_bug_report(test_name, exception, screenshot_path):
    details = {
        'title': f"Bug in {test_name}",
        'date': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'reporter': 'Automated Test',
        'description': f"An exception occurred: {str(exception)}",
        'step1': 'Open the login page.',
        'step2': 'Enter the username and password.',
        'step3': 'Click the login button.',
        'expected_result': 'User should be logged in successfully.',
        'actual_result': 'An error occurred during the login process.',
        'severity': 'High',
        'priority': 'High'
    }

    template = BugReportTemplate()

    bug_report_name = f"{test_name}_bug_report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    bug_report_path = os.path.join(os.path.dirname(__file__), 'bug_reports', bug_report_name)
    os.makedirs(os.path.dirname(bug_report_path), exist_ok=True)

    template.generate_report(details, screenshot_path, bug_report_path)

    print(f"Bug report generated: {bug_report_path}")
    send_email(bug_report_path, 'receiver@gmail.com')


def test_login_valid_credentials(driver):
    test_name = 'test_login_valid_credentials'
    try:
        login(driver, 'mail123@gmail.com', '1234563')
        WebDriverWait(driver, 10).until(EC.url_changes('https://tutorialsninja.com/demo/index.php?route=account/login'))
        assert driver.current_url == 'https://tutorialsninja.com/demo/index.php?route=account/account'
    except Exception as e:
        screenshot_path = capture_screenshot(driver, test_name)
        generate_bug_report(test_name, e, screenshot_path)
        raise e


def test_login_invalid_email(driver):
    test_name = 'test_login_invalid_username'
    try:
        login(driver, 'invalid_email', '123456')
        error_message = driver.find_element(By.CSS_SELECTOR, ".alert-dismissible")
        assert 'Warning: No match for E-Mail Address and/or Password.' in error_message.text
    except Exception as e:
        screenshot_path = capture_screenshot(driver, test_name)
        generate_bug_report(test_name, e, screenshot_path)
        raise e


def test_login_invalid_password(driver):
    test_name = 'test_login_invalid_password'
    try:
        login(driver, 'mail123@gmail.com', '12345622')
        error_message = driver.find_element(By.CSS_SELECTOR, ".alert-dismissible")
        assert 'Warning: No match for E-Mail Address and/or Password.' in error_message.text
    except Exception as e:
        screenshot_path = capture_screenshot(driver, test_name)
        generate_bug_report(test_name, e, screenshot_path)
        raise e


def test_login_empty_username(driver):
    test_name = 'test_login_empty_username'
    try:
        login(driver, '', '123456')
        error_message = driver.find_element(By.CSS_SELECTOR, ".alert-dismissible")
        assert 'Warning: No match for E-Mail Address and/or Password.' in error_message.text
    except Exception as e:
        screenshot_path = capture_screenshot(driver, test_name)
        generate_bug_report(test_name, e, screenshot_path)
        raise e


def test_login_empty_password(driver):
    test_name = 'test_login_empty_password'
    try:
        login(driver, 'mail123@gmail.com', '')
        error_message = driver.find_element(By.CSS_SELECTOR, ".alert-dismissible")
        assert 'Warning: No match for E-Mail Address and/or Password.' in error_message.text
    except Exception as e:
        screenshot_path = capture_screenshot(driver, test_name)
        generate_bug_report(test_name, e, screenshot_path)
        raise e
